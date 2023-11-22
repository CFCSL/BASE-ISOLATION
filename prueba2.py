# -*- coding: utf-8 -*-
"""
Created on Mon Nov 20 18:01:09 2023

@author: cfcpc2
"""
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement

def replace_figure(doc, old_figure_path, new_figure_path, old_caption, new_caption):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if old_figure_path in run.text:
                # Replace the figure
                run.clear()
                run.add_picture(new_figure_path, width=Inches(3.25))

                # Replace the caption text
                if old_caption in paragraph.text:
                    paragraph.clear()
                    paragraph.add_run().add_text(new_caption)

# Load the existing document
document = Document('demo.docx')

# Path to the new figure you want to replace the old one with
new_figure_path = 'new_figure.png'

# Path to the old figure you want to replace
old_figure_path = 'monty-truth.png'

# Caption to replace
old_caption = 'Figure 1: Monty Truth'

# New caption
new_caption = 'Figure 1: New Figure'

# Replace the figure and caption
replace_figure(document, old_figure_path, new_figure_path, old_caption, new_caption)

# Save the modified document
document.save('modified_demo2.docx')