# -*- coding: utf-8 -*-
"""
Created on Mon Nov 20 17:19:10 2023

@author: cfcpc2
"""

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement


document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True




document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_heading('Header 2', level=2)
document.add_paragraph('Intense quote', style='Intense Quote')
document.add_paragraph(
	'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
	'first item in ordered list', style='List Number'
)

document.add_picture('monty-truth.png', width=Inches(3.25))
# Add a caption to the picture
document.paragraphs[-1].add_run().add_text('Figure 1: Monty Truth')

records = (
	(3, '101', 'Spam'),
	(7, '422', 'Eggs'),
	(4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
	row_cells = table.add_row().cells
	row_cells[0].text = str(qty)
	row_cells[1].text = id
	row_cells[2].text = desc
	


document.add_page_break()

document.save('demo.docx')

def replace_figure(doc, old_figure_path, new_figure_path, old_caption, new_caption):
	for paragraph in doc.paragraphs:
		for run in paragraph.runs:
			if old_figure_path in run.text:
				# Replace the figure
				run.clear()
				run.add_picture(new_figure_path, width=Inches(3.25))

				# Replace the caption text
				if old_caption in paragraph.text:
					paragraph.clear()
					paragraph.add_run().add_text(new_caption)

# Load the existing document
document = Document('demo.docx')

# Path to the new figure you want to replace the old one with
new_figure_path = 'new_figure.png'

# Path to the old figure you want to replace
old_figure_path = 'monty-truth.png'

# Caption to replace
old_caption = 'Figure 1: Monty Truth'

# New caption
new_caption = 'Figure 1: New Figure'

# Replace the figure and caption
replace_figure(document, old_figure_path, new_figure_path, old_caption, new_caption)

# Save the modified document
document.save('modified_demo2.docx')